"""Extract BSG docx into plain text."""
from docx import Document
import sys

src = "notes/BSG 2025 UPDATED BRANDED STANDARD GUIDELINES 2025.docx"
dst = "notes/pdf_text/bsg_2025.txt"

d = Document(src)
lines = []
for p in d.paragraphs:
    style = p.style.name if p.style else ""
    text = p.text.strip()
    if not text:
        lines.append("")
        continue
    if style.startswith("Heading"):
        lines.append(f"\n## [{style}] {text}\n")
    else:
        lines.append(text)

# Also extract tables
for ti, table in enumerate(d.tables):
    lines.append(f"\n--- TABLE {ti} ---")
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
        lines.append(" || ".join(cells))

with open(dst, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))

print(f"Wrote {len(lines)} lines to {dst}")
